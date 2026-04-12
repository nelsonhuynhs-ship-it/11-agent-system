"""Programmatic python-docx template for the weekly market report.

Builds the 4C structure with Vietnamese headers:
  I. COSTING — top carriers per lane from parquet
  II. CAPACITY — team-entered scores
  III. CHALLENGE & CHANCE — ranked catalysts
  IV. FORECAST TUẦN (W+1) — low/base/high scenarios
  V. BACKTEST TUẦN (W) — previous-week accuracy
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from ..schemas import (
    CapacitySignal,
    Catalyst,
    CostingItem,
    ForecastScenario,
)


def build_report(
    prev_week: str,
    next_week: str,
    costing: list[CostingItem],
    capacity: list[CapacitySignal],
    catalysts: list[Catalyst],
    forecast: list[ForecastScenario],
    backtest_rows: list[dict],
    output_path: Path,
    model_version: str = "baseline-v1",
) -> Path:
    """Render the 4C report DOCX. Returns the output path."""
    doc = Document()

    _title(doc, f"BÁO CÁO THỊ TRƯỜNG TUẦN {prev_week} & DỰ ĐOÁN TUẦN {next_week}")
    _meta(doc, f"Generated: {datetime.now():%Y-%m-%d %H:%M} | Model: {model_version}")

    _section_costing(doc, costing)
    _section_capacity(doc, capacity)
    _section_catalysts(doc, catalysts)
    _section_forecast(doc, next_week, forecast)
    _section_backtest(doc, prev_week, backtest_rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ── Title & metadata ──────────────────────────────────────────────────────────

def _title(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)


def _meta(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)


# ── Section I. Costing ─────────────────────────────────────────────────────────

def _section_costing(doc, items: list[CostingItem]) -> None:
    _heading(doc, "I. COSTING")
    if not items:
        doc.add_paragraph("(Không có dữ liệu costing từ parquet cho tuần này.)")
        return
    doc.add_paragraph("Các giá tốt mà Pudong đang có:")
    for lane in ("WC", "EC", "GULF"):
        lane_items = [i for i in items if i.lane == lane]
        if not lane_items:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{lane}:")
        run.bold = True
        for it in lane_items:
            valid = ""
            if it.valid_to:
                valid = f" (valid to {it.valid_to.strftime('%d %b')})"
            spread = ""
            if it.spread_vs_lane_avg:
                sign = "-" if it.spread_vs_lane_avg < 0 else "+"
                spread = f" [{sign}${abs(it.spread_vs_lane_avg):.0f} vs avg]"
            doc.add_paragraph(
                f"{it.carrier} {it.rate_type}: "
                f"${it.price:.0f}/{it.container}{valid}{spread}",
                style="List Bullet",
            )


# ── Section II. Capacity ──────────────────────────────────────────────────────

def _section_capacity(doc, signals: list[CapacitySignal]) -> None:
    _heading(doc, "II. CAPACITY")
    if not signals:
        doc.add_paragraph("(Chưa có input capacity từ CS team cho tuần này.)")
        return
    avg_score = sum(s.score for s in signals) / len(signals)
    doc.add_paragraph(
        f"Tổng hợp từ CS team: {len(signals)} entries — "
        f"điểm trung bình {avg_score:.1f}/5"
    )
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Carrier"
    hdr[1].text = "Lane"
    hdr[2].text = "Status"
    hdr[3].text = "Score"
    hdr[4].text = "Notes"
    for s in signals:
        row = table.add_row().cells
        row[0].text = s.carrier
        row[1].text = s.lane
        row[2].text = s.status
        row[3].text = f"{s.score}/5"
        row[4].text = s.notes or ""


# ── Section III. Challenge & Chance ───────────────────────────────────────────

def _section_catalysts(doc, catalysts: list[Catalyst]) -> None:
    _heading(doc, "III. CHALLENGE & CHANCE")
    if not catalysts:
        doc.add_paragraph("(Chưa có catalyst nào được thu thập cho tuần này.)")
        return
    doc.add_paragraph(f"{len(catalysts)} catalysts — ranked theo impact magnitude")
    for c in catalysts:
        lanes = ",".join(c.affected_lanes) if c.affected_lanes else "ALL"
        eff = c.effective_date.isoformat() if c.effective_date else "n/a"
        title_p = doc.add_paragraph(style="List Bullet")
        bold = title_p.add_run(f"[{c.impact_magnitude}] {c.headline}")
        bold.bold = True
        title_p.add_run(f"  — source: {c.source}, eff: {eff}")
        doc.add_paragraph(c.body)
        doc.add_paragraph(
            f"→ Impact: {c.impact_direction} {c.impact_magnitude} on {lanes} "
            f"(confidence {c.confidence:.0%})"
        )


# ── Section IV. Forecast ──────────────────────────────────────────────────────

def _section_forecast(
    doc,
    next_week: str,
    scenarios: list[ForecastScenario],
) -> None:
    _heading(doc, f"IV. FORECAST TUẦN {next_week}")
    if not scenarios:
        doc.add_paragraph("(Chưa có forecast cho tuần tới.)")
        return
    for f in scenarios:
        p = doc.add_paragraph(style="List Bullet")
        bold = p.add_run(f"{f.lane} {f.container}: ")
        bold.bold = True
        p.add_run(
            f"base ${f.base_case:.0f} "
            f"(range ${f.low_case:.0f}–${f.high_case:.0f}) "
            f"confidence {f.confidence:.0%}"
        )
        if f.rationale:
            doc.add_paragraph(f"   Rationale: {f.rationale}")
        if f.trigger_catalyst_ids:
            doc.add_paragraph(
                "   Triggers: " + ", ".join(f.trigger_catalyst_ids)
            )


# ── Section V. Backtest ───────────────────────────────────────────────────────

def _section_backtest(doc, prev_week: str, rows: list[dict]) -> None:
    _heading(doc, f"V. BACKTEST TUẦN {prev_week}")
    if not rows:
        doc.add_paragraph(
            "(Không có forecast tuần trước để backtest, hoặc parquet thiếu dữ liệu actual.)"
        )
        return
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Lane"
    hdr[1].text = "Forecast"
    hdr[2].text = "Actual"
    hdr[3].text = "Error $"
    hdr[4].text = "Error %"
    errors = []
    for r in rows:
        row = table.add_row().cells
        row[0].text = str(r.get("lane", ""))
        row[1].text = f"${float(r.get('forecast_base', 0)):.0f}"
        row[2].text = f"${float(r.get('actual_avg', 0)):.0f}"
        row[3].text = f"${float(r.get('error_abs', 0)):+.0f}"
        row[4].text = f"{float(r.get('error_pct', 0)):+.1f}%"
        try:
            errors.append(abs(float(r.get("error_pct", 0))))
        except (TypeError, ValueError):
            pass
    if errors:
        avg_err = sum(errors) / len(errors)
        rating = _accuracy_rating(avg_err)
        doc.add_paragraph(
            f"Week accuracy score: {avg_err:.1f}% ({rating})"
        )


def _accuracy_rating(avg_err_pct: float) -> str:
    if avg_err_pct < 5:
        return "excellent"
    if avg_err_pct < 10:
        return "good"
    if avg_err_pct < 20:
        return "fair"
    return "poor"
